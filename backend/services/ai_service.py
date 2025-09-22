# backend/services/ai_service.py
import os
from typing import Dict, Any
from datetime import datetime
from docx import Document

class NDISAIService:
    """Service for generating NDIS documents synchronously"""

    def generate_ndis_document(
        self, 
        template: Dict, 
        client: Dict, 
        coordinator: Dict, 
        context: str, 
        uploaded_files: list
    ) -> Dict[str, Any]:
        """Generate a real DOCX document"""

        template_name = template['name']
        client_name = client['full_name']

        # Ensure output folder exists
        output_dir = os.path.join("output", "generated")
        os.makedirs(output_dir, exist_ok=True)

        # File name
        timestamp = datetime.now().strftime("%Y%m%d%H%M")
        file_name = f"{template_name}_{client_name}_{timestamp}.docx"
        output_path = os.path.join(output_dir, file_name)

        # Create DOCX
        doc = Document()
        doc.add_heading(template_name, level=1)

        # Participant info
        doc.add_heading("Participant Information", level=2)
        doc.add_paragraph(f"Name: {client['full_name']}")
        doc.add_paragraph(f"NDIS Number: {client['ndis_number']}")
        doc.add_paragraph(f"Plan Period: {client.get('plan_start_date', 'N/A')} → {client.get('plan_end_date', 'N/A')}")

        # Coordinator info
        doc.add_heading("Support Coordinator Information", level=2)
        doc.add_paragraph(f"Coordinator: {coordinator.get('coordinator_name', 'Unknown')}")
        doc.add_paragraph(f"Provider: {coordinator.get('provider_name', 'Unknown')}")
        doc.add_paragraph(f"Email: {coordinator.get('email', 'N/A')}")
        doc.add_paragraph(f"Phone: {coordinator.get('phone', 'N/A')}")

        # Context
        doc.add_heading("Reason / Context", level=2)
        doc.add_paragraph(context)

        # Conditions
        doc.add_heading("Conditions", level=2)
        conditions = client.get('conditions', [])
        if conditions:
            for cond in conditions:
                doc.add_paragraph(f"• {cond}")
        else:
            doc.add_paragraph("No conditions recorded")

        # Goals
        doc.add_heading("Goals", level=2)
        goals = []
        if client.get('goals_advocacy'): goals.append("Advocacy")
        if client.get('goals_support'): goals.append("Support")
        if client.get('goals_capacity_building'): goals.append("Capacity Building")
        if client.get('goals_funding_increase'): goals.append("Funding Increase")
        if client.get('goals_other'): goals.append(client['goals_other'])
        if goals:
            for g in goals:
                doc.add_paragraph(f"• {g}")
        else:
            doc.add_paragraph("No goals provided")

        # Save DOCX
        doc.save(output_path)

        return {
            "document_name": file_name,
            "content": f"{template_name} generated for {client_name}",
            "output_path": output_path
        }
